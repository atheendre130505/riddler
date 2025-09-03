import os
import PyPDF2
from docx import Document
from typing import Dict, Optional
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor:
    """
    Handles file processing for various document types
    Supports PDF, DOCX, and TXT files
    """
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt'}
    
    def process_file(self, file_path: str) -> Dict:
        """
        Process a file and extract text content
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Dict containing extracted content and metadata
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Extract text based on file type
            if file_extension == '.pdf':
                content = self.extract_text_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                content = self.extract_text_docx(file_path)
            elif file_extension == '.txt':
                content = self.extract_text_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Clean and validate content
            content = self.clean_content(content)
            
            if not content or len(content.strip()) < 10:
                raise ValueError("No meaningful content extracted from file")
            
            return {
                "content": content,
                "file_type": file_extension,
                "file_size": os.path.getsize(file_path),
                "word_count": len(content.split()),
                "char_count": len(content)
            }
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise
    
    def extract_text_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_text_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_text_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                content = file.read()
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")
    
    def clean_content(self, content: str) -> str:
        """
        Clean and normalize extracted content
        
        Args:
            content: Raw extracted content
            
        Returns:
            Cleaned content
        """
        if not content:
            return ""
        
        # Remove excessive whitespace
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Only keep non-empty lines
                cleaned_lines.append(line)
        
        # Join lines with single newlines
        cleaned_content = '\n'.join(cleaned_lines)
        
        # Remove excessive spaces
        import re
        cleaned_content = re.sub(r'\s+', ' ', cleaned_content)
        
        return cleaned_content.strip()
    
    def get_file_info(self, file_path: str) -> Dict:
        """
        Get basic file information
        
        Args:
            file_path: Path to file
            
        Returns:
            Dict with file metadata
        """
        try:
            stat = os.stat(file_path)
            return {
                "filename": os.path.basename(file_path),
                "size": stat.st_size,
                "extension": os.path.splitext(file_path)[1].lower(),
                "modified": stat.st_mtime
            }
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {str(e)}")
            return {}


